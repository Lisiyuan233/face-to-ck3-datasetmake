from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "CK3_FACE_TO_DNA_TRAINING_PLAN.md"
OUTPUT = ROOT / "docs" / "CK3_FACE_TO_DNA_TRAINING_PLAN.docx"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5E6A78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "C9D3DF"


def set_run_font(run, size=None, color=None, bold=None, italic=None, mono=False):
    ascii_font = "Consolas" if mono else "Calibri"
    east_asia_font = "Microsoft YaHei"
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color="000000", bold=False):
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, color=BORDER, size=5):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def choose_widths(rows):
    cols = len(rows[0])
    if cols == 2:
        return [2700, 6660]
    if cols == 3:
        return [1800, 1900, 5660]
    if cols == 4:
        return [1350, 3100, 1800, 3110]
    base = PAGE_WIDTH_DXA // cols
    widths = [base] * cols
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def add_numbering_definition(doc, fmt):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), fmt)
    lvl.append(numfmt)
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    lvl.append(lvltext)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text, base_size=11, color="000000"):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=max(8.8, base_size - 0.5), color=DARK_BLUE, mono=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, color=color)


def clean_table_text(text):
    text = text.strip()
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.line_spacing = 1.20
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_FILL)
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), BLUE)
    left.set(qn("w:space"), "8")
    pbdr.append(left)
    ppr.append(pbdr)
    add_inline(p, text, base_size=10.5, color=INK)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F4F7")
    ppr.append(shd)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=8.5, color=INK, mono=True)


def add_markdown_table(doc, rows):
    widths = choose_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            if r_idx == 0:
                shade_cell(cell, HEADER_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if c_idx == 0 and len(rows[0]) <= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif len(clean_table_text(text)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, clean_table_text(text), base_size=8.8 if len(rows[0]) >= 4 else 9.2, color=INK)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    after.paragraph_format.line_spacing = 0.4


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, text, fld_char_end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    set_style_font(title, 26, INK, True)
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, 11, MUTED, False)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        add_inline(hp, "FACE TO CK3 DNA  |  完整训练方案", base_size=8.5, color=MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        prefix = fp.add_run("2026-07-21  ·  ")
        set_run_font(prefix, size=8.5, color=MUTED)
        add_field(fp, "PAGE")


def parse_markdown(doc, text):
    bullet_num_id = add_numbering_definition(doc, "bullet")
    decimal_num_id = None
    in_ordered_list = False
    lines = text.splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if not line:
            in_ordered_list = False
            i += 1
            continue

        if line.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            add_code_block(doc, code)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1]):
            rows = []
            header = [x.strip() for x in line.strip("|").split("|")]
            rows.append(header)
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([x.strip() for x in lines[i].strip("|").split("|")])
                i += 1
            add_markdown_table(doc, rows)
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_inline(p, line[2:].strip(), base_size=26, color=INK)
            first_title = False
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:].strip(), base_size=16, color=BLUE)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:].strip(), base_size=13, color=BLUE)
        elif line.startswith("> "):
            if not first_title and "版本 v" in line:
                p = doc.add_paragraph(style="Subtitle")
                add_inline(p, line[2:].strip(), base_size=11, color=MUTED)
            else:
                add_callout(doc, line[2:].strip())
        elif re.match(r"^\s*-\s+", line):
            p = doc.add_paragraph()
            apply_num(p, bullet_num_id)
            add_inline(p, re.sub(r"^\s*-\s+", "", line), base_size=11)
        elif re.match(r"^\s*\d+\.\s+", line):
            if not in_ordered_list or decimal_num_id is None:
                decimal_num_id = add_numbering_definition(doc, "decimal")
            in_ordered_list = True
            p = doc.add_paragraph()
            apply_num(p, decimal_num_id)
            add_inline(p, re.sub(r"^\s*\d+\.\s+", "", line), base_size=11)
        else:
            in_ordered_list = False
            p = doc.add_paragraph()
            add_inline(p, line, base_size=11)
        i += 1


def audit_document(doc):
    assert len(doc.sections) == 1
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.50
    assert round(section.left_margin.inches, 2) == 1.00
    for table in doc.tables:
        tblpr = table._tbl.tblPr
        tblw = tblpr.find(qn("w:tblW"))
        tblind = tblpr.find(qn("w:tblInd"))
        assert tblw is not None and int(tblw.get(qn("w:w"))) == PAGE_WIDTH_DXA
        assert tblind is not None and int(tblind.get(qn("w:w"))) == TABLE_INDENT_DXA
        grid_widths = [int(x.get(qn("w:w"))) for x in table._tbl.tblGrid]
        assert sum(grid_widths) == PAGE_WIDTH_DXA
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tcw is not None and int(tcw.get(qn("w:w"))) == grid_widths[idx]


def main():
    doc = Document()
    configure_document(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.core_properties.title = "Face to CK3 DNA 完整训练方案"
    doc.core_properties.subject = "510,000 组 CK3 男性肖像到 DNA 的训练、评估与部署规划"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "CK3, DNA, face, multi-task learning, training plan"
    audit_document(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
